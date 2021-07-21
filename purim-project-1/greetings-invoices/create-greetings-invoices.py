import datetime
import xlrd
# https://www.geeksforgeeks.org/reading-excel-file-using-python/
import docx
# https://stackabuse.com/reading-and-writing-ms-word-files-in-python-via-python-docx-module/
# https://pbpython.com/python-word-template.html
# https://python-docx.readthedocs.io/en/latest/user/quickstart.html#adding-a-table
# CHANGE FONT AND SIZE
# https://python-docx.readthedocs.io/en/latest/user/text.html#paragraph-properties
import os
# OS https://stackabuse.com/creating-and-deleting-directories-with-python/
import shutil
from docx.shared import Pt
import logging

logging.basicConfig(filename='output/error.txt', encoding='utf-8', level=logging.DEBUG)

GREETING_COST = 5
FREE_RECIP_GREETINGS = 10
RECIPROCITY_COST = 25
UNLIMITED_COST = 250
COLUMNS = 3
SHOW_ID = False


class Member:
  	def __init__(self, id, firstName, lastName):
	    self.id = id
	    self.firstName = firstName
	    self.lastName = lastName
	    self.cellText = str(id) + " " + firstName + " " + lastName
	    self.greeters = []


class Order: 
	def __init__(self, senderID, amountPaid, reciprocity, receivers):
	    self.senderID = int(senderID)
	    self.amountPaid = int(amountPaid)
	    self.reciprocity = reciprocity
	    self.receivers = receivers
	    self.initialCost = len(receivers) * GREETING_COST
	    self.reciprocityCost = reciprocity * RECIPROCITY_COST
	    self.reciprocityOverCost = 0
	    self.amountDue = 0
	    self.reciprocals = []

	def calculateAmountDue(self):
		if self.amountPaid == UNLIMITED_COST:
			self.amountDue = 0
		else: self.amountDue = (	 
			self.initialCost + 
			self.reciprocityCost + 
			self.reciprocityOverCost -
			self.amountPaid
		)


class Greeting:
  	def __init__(self, id):
	    self.recipient_id = id
	    self.senders = []


# READ MEMBER LIST, CREATE MEMBER OBJECTS, APPEND TO LIST  
def readMemberList(members):
	# OPEN EXCEL WORKBOOK, SELECT SHEET, CREATE DOCX OBJECT
	members_file = xlrd.open_workbook("./members.xlsx") 
	sheet = members_file.sheet_by_index(0) 

	numRows = sheet.nrows
	numCols = sheet.ncols

	for i in range(1, numRows):
		id = int(sheet.cell_value(i, 0))
		firstName = sheet.cell_value(i, 2).strip()
		lastName = sheet.cell_value(i, 1).strip()
		members.append(Member(id, firstName, lastName))


# WRITE MEMBER NAMES/ID TO GREETING REQUEST WORD DOC
def printMemberNames(members):

	doc = docx.Document()

	rowsWordDoc = int(len(members) / COLUMNS + 1)
	table = doc.add_table(rows=rowsWordDoc, cols=COLUMNS)

	j = 0
	for i in range(0,len(members), COLUMNS):
		table.cell(j, 0).text = members[i].cellText
		if(i+1 < len(members)):
			table.cell(j, 1).text = members[i+1].cellText
		if(i+2 < len(members)):
			table.cell(j, 2).text = members[i+2].cellText
		j+=1

	doc.save("./output/" +'member_names.docx')


# READ ORDER LIST, CREATE ORDER OBJECTS, APPEND TO LIST  
def readOrderList(orders):
	# OPEN EXCEL WORKBOOK, SELECT SHEET, CREATE DOCX OBJECT
	orders_file = xlrd.open_workbook("./orders.xlsx") 
	sheet = orders_file.sheet_by_index(0)

	numRows = sheet.nrows
	numCols = sheet.ncols

	for i in range(1, numRows):
		senderID = int(sheet.cell_value(i, 0))
		reciprocity = True if sheet.cell_value(i, 1) else False
		amountPaid = sheet.cell_value(i, 2)
		receivers_string = sheet.cell_value(i, 3)

		# IF ONLY SINGLE NUMBER IN CELL, RECOGNIZED AS FLOAT NOT STRING
		if type(receivers_string) is float:
			recievers_list = int(receivers_string)
		else: 
			string_list = receivers_string.split()
			map_object = map(int, string_list)
			receivers_list = list(map_object)

		for j in receivers_list:
			if len(members) > j - 1:
				members[j-1].greeters.append(senderID)
			else: 
				logging.error("In row: " + str(i+1) + "\tinvalid member ID given: " + str(j))
		orders.append(Order(senderID, amountPaid, reciprocity, receivers_list))


# CALCULATE RECIPROCITY GREETINGS (RECIPROCALS) TO BE SENT PER SENDER
def calculateReciprocity(orders):
	for myOrder in orders:
		if myOrder.reciprocity == True:
			for otherOrder in orders: 
				if myOrder.senderID in otherOrder.receivers: 
					if otherOrder.senderID not in myOrder.receivers:
						myOrder.reciprocals.append(otherOrder.senderID)

	for order in orders:
		numRecips = len(order.reciprocals)

		if numRecips > FREE_RECIP_GREETINGS:
			overCost = (numRecips - FREE_RECIP_GREETINGS) * GREETING_COST
			order.reciprocityOverCost = overCost


# PRINT TO CONSOLE OR TEXT FILE ALL ORDER DETAILS AND BALANCES
def printOrderDetails(orders):

	orders_file = open("./output/all_orders.txt","w+") 

	for i in orders:
		i.calculateAmountDue()
		orders_file.write(
			"Sender Name: " + "\t" +  
			members[i.senderID-1].firstName + " " + 
			members[i.senderID-1].lastName+ "\n"+
			"Sender ID:   " + "\t" + str(i.senderID) + "\n" +
			"Reciprocity: " + "\t" + str(i.reciprocity)+ "\n"+
			"Amount Paid: " + "\t" + str(i.amountPaid)+ "\n"+
			"Initial Cost:" + "\t" + str(i.initialCost)+ "\n"+
			"Recip. Cost: " + "\t" + str(i.reciprocityCost)+ "\n"+
			"Over cost:   " + "\t" + str(i.reciprocityOverCost)+ "\n"+
			"Recipients:  " + "\t" + str(i.receivers)+ "\n"+
			"Reciprocals: " + "\t" + str(i.reciprocals)+ "\n"+
			"TOTAL DUE:   " + "\t" + str(i.amountDue)+ " NIS\n\n"
		)

		if i.amountDue > 0:

			invoices_file = open("./invoices/invoice_" + str(i.senderID) + ".txt","w+") 

			invoices_file.write(
				"--------------------------------------------------------------------\n\n" +
				"Name: " + "\t" +  
				members[i.senderID-1].firstName + " " + 
				members[i.senderID-1].lastName+ "\n\n" +
				"\tNumber of greetings:        \t   " + str(len(i.receivers)) + "\n" +
				"X\tCost per greeting:         \t nis " + str(GREETING_COST) + "\n" +
				"=\tGreeting cost:             \t nis " + str(i.initialCost) + "\n\n" + 
				"\tReciprocity option cost:    \t nis " + str(i.reciprocityCost) + "\n\n" +
				"\t# of reciprocity greetings: \t   " + str(len(i.reciprocals)) + "\n" +
				"X\tGreeting cost (after 10):  \t nis " + str(GREETING_COST) + "\n" +
				"=\tReciprocity greeting cost: \t nis " + str(i.reciprocityOverCost) + "\n\n" + 

				"Total cost:   \t nis " + str(i.initialCost + i.reciprocityCost + 
					i.reciprocityOverCost)+ "\n" +

				"Amount Paid:  \t nis " + str(i.amountPaid) + "\n" +
			
				"TOTAL DUE:    \t nis " + str(i.amountDue)+ "\n\n"

				"--------------------------------------------------------------------\n\n" 

			)


# PRINT INVOICES OF ANY OUTSTANDING BALANCES
def printOrderInvoices(orders):

	doc = docx.Document()
	orders_file = open("./output/all_orders.txt","w+") 

	for i in orders:
	
		i.calculateAmountDue()
	
		orders_file.write(
			"Sender Name: " + "\t" +  
			members[i.senderID-1].firstName + " " + 
			members[i.senderID-1].lastName+ "\n"+
			"Sender ID:   " + "\t" + str(i.senderID) + "\n" +
			"Reciprocity: " + "\t" + str(i.reciprocity)+ "\n"+
			"Amount Paid: " + "\t" + str(i.amountPaid)+ "\n"+
			"Initial Cost:" + "\t" + str(i.initialCost)+ "\n"+
			"Recip. Cost: " + "\t" + str(i.reciprocityCost)+ "\n"+
			"Over cost:   " + "\t" + str(i.reciprocityOverCost)+ "\n"+
			"Recipients:  " + "\t" + str(i.receivers)+ "\n"+
			"Reciprocals: " + "\t" + str(i.reciprocals)+ "\n"+
			"TOTAL DUE:   " + "\t" + str(i.amountDue)+ " NIS\n\n"
		)

		if i.amountDue > 0:

			run1 = doc.add_paragraph().add_run(
				"- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - " + 
				"- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -\n\n" +
				"Chug Na'avah Tehilla Mishloach Manot\n\n" +
				"INVOICE\n\n" +
				"Name: " + "\t" +  
				members[i.senderID-1].firstName + " " + 
				members[i.senderID-1].lastName+ "\n\n" +
				" \tNumber of greetings:       \t\t" + str(len(i.receivers)) + "\n" +
				"X\tCost per greeting:         \t\t" + str(GREETING_COST) + " nis\n" +
				"=\tGreeting cost:             \t\t" + str(i.initialCost) + " nis\n\n" + 
				" \tReciprocity option cost:   \t\t" + str(i.reciprocityCost) + " nis\n\n" +
				"\t# of reciprocity greetings: \t\t" + str(len(i.reciprocals)) + "\n" +
				"X\tCost per greeting after 10:\t" + str(i.reciprocityCost) + " nis\n" +
				"=\tReciprocity greeting cost: \t\t" + str(i.reciprocityOverCost) + " nis\n\n" + 

				"\tTotal cost:   \t\t\t\t" + str(i.initialCost + i.reciprocityCost + 
					i.reciprocityOverCost)+ " nis\n" +

				"\tAmount Paid:  \t\t\t" + str(i.amountPaid) + " nis\n" +
			
				"\tAMOUNT DUE:   \t\t" + str(i.amountDue)+ " nis\n\n" +

				"Please send a shekel check for the balance due to Emunah. " + 
				"Please mark the envelope N.T. and enclose this invoice. " + 
				"Thank you for participating in this project!\n\n"

				"- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - " + 
				"- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - " 

			)

			run1.font.size = Pt(14)
			run1.font.name = "Goudy Old Style"

			doc.add_page_break()
			
	doc.save("./output/" + "all_invoices" + '.docx')


# PRINT WORD DOC GREETINGS IN SEPARATE FOLDER IN CURRENT DIRECTORY
def printGreetings():

	thisyear = datetime.datetime.now().year

	doc = docx.Document()

	for i in members:
		if len(i.greeters) > 0:

			doc.add_picture(os.getcwd() + "/purim.jpg", 
				width=docx.shared.Inches(6), height=docx.shared.Inches(2))

			showID = "(" + str(i.id) + ")" if SHOW_ID else ""

			run1 = doc.add_paragraph().add_run(
				"\n" + str(3760 + thisyear) + " Purim " + str(thisyear) + "\n\n" + 
				"Dear " + i.firstName + " " + i.lastName + ", " + showID + 
				"\n\nIn lieu of Mishloach Manot, a donation has been made to " +
				"Na'avah Tehilla Emunah in your honor by your friends listed below:\n"
				)

			run1.font.size = Pt(16)
			run1.font.name = "Goudy Old Style"

			# for j in i.greeters:
			# 	show_id = "(" + str(j) + ")" if SHOW_ID else ""

			sent_greetings = len(i.greeters)
			
			table = doc.add_table(rows=12, cols=3)

			k = 0
			for h in range(0,sent_greetings, 3):
				table.cell(k, 0).text = members[h].firstName + ' ' + members[h].lastName
				if(h+1 < sent_greetings):
					table.cell(k, 1).text = members[h+1].firstName + ' ' + members[h+1].lastName
				if(h+2 < sent_greetings):
					table.cell(k, 2).text = members[h+2].firstName + ' ' + members[h+2].lastName
				k+=1

				# run2 = doc.add_paragraph().add_run(
				# 	members[j-1].firstName + " " + 
				# 	members[j-1].lastName + " " + show_id
				# )
				# run2.font.size = Pt(12)
				# run2.font.name = "Goudy Old Style"

			doc.add_page_break()
			
	doc.save("./output/" + "all_greetings" + '.docx')


# OUTPUT LIST OF ALL RECIPIENTS
def printAllRecipients():
	recipients_file = open("./output/all_recipients.txt","w+") 

	for i in members:
		if len(i.greeters) > 0:
			recipients_file.write(i.firstName + " " + i.lastName + "\n")


# CREATE DIRECTORY FOR ALL OUTPUT FILES
def createFolder():
	path = os.getcwd() + "/output"
	
	try:
		os.mkdir(path)
	except OSError:
		print('Output folder already exists.')


members = []
orders = []
greetings = []


createFolder()
readMemberList(members)
# printMemberNames(members)
readOrderList(orders)
calculateReciprocity(orders)
printOrderInvoices(orders)
printGreetings()
printAllRecipients()


def originalPrint():
	print(" Sender ID/Name:", i.senderID, 
			members[i.senderID-1].firstName, 
			members[i.senderID-1].lastName, "\n", 
			"Reciprocity:", "\t", i.reciprocity, "\n", 
			"Amount Paid:", "\t", i.amountPaid, "\n", 
			"Initial Cost:", "\t", i.initialCost, "\n", 
			"Recip. Cost:", "\t", i.reciprocityCost, "\n", 
			"Over cost:", "\t", i.reciprocityOverCost, "\n",
			"Amount due:", "\t", i.amountDue, "\n",
			"Recipients:", "\t", i.receivers, "\n",
			"Reciprocals:", "\t", i.reciprocals, "\n")
